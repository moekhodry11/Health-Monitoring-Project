import schedule
import time
import pandas as pd
import telebot
import matplotlib.pyplot as plt
from docx.shared import Inches
import matplotlib.dates as mdates
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docx.shared import Cm

# FOR GENERATING MANY WORDS
counter = 1

def extract_data_and_generate_report():
    global counter
    # read the daata sheet from excel file
    df = pd.read_excel('data/edited_data .xlsx', sheet_name='daata')
    # read the 'Measurement', 'Average', 'Max', 'Min' columns from our dataframe
    important_data = df[['Measurement', 'Average', 'Max', 'Min']]

    # Create a new Word document
    doc_name = f'reports/report{counter}.docx'
    doc_name_print = f'report{counter}.docx'
    counter += 1
    doc = Document()

    # for styling the document file
    # Add a heading to the document
    doc.add_heading('Important Data', level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # for plotting BPM vs time
    # specify x and y  axises
    x_column = 'Time'
    y_column = 'BPM'

    # To adjust the time format as H:M:S
    df[x_column] = pd.to_datetime(df[x_column], format='%H:%M:%S')

    # create the plot   BPM AND TIME
    plt.figure(figsize=(6, 4))
    fig, ax = plt.subplots()
    ax.plot(df['Index'], df[y_column])
    ax.set_xlabel('spontinous reading')
    ax.set_ylabel('BPM')
    ax.set_title('BPM')

    # format the time axis label to appear as H:M:S
    # plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%H:%M:%S'))
    # plt.gcf().autofmt_xdate()

    # save the plot as a png
    plot_filename = 'BPM_plot.png'
    plt.savefig(plot_filename)

    # create a hisogram for SPO2
    plt.figure(figsize=(6, 4))
    plt.hist(df['SPO2'], bins=10)
    plt.xlabel('SPO2')
    plt.ylabel('Frequency')
    plt.title('Histogram')

    # save the plot as a png
    hist_filename = "histogram.png"
    plt.savefig(hist_filename)

    # Convert the important data to a table in Word
    table = doc.add_table(1, len(important_data.columns))
    # style the heading of the table
    table.style = 'Light List Accent 1'
    # change the table dimensions
    new_width = Cm(10)
    change_table_dimensions(table, new_width)

    # Add the header rows.
    for j in range(len(important_data.columns)):
        table.cell(0, j).text = important_data.columns[j]

    # fill the table with the data
    for i in range(3):
        # Add a row to the table
        row_cells = table.add_row().cells
        # fill the data
        for j in range(len(important_data.columns)):
            row_cells[j].text = str(important_data.values[i, j])


    # add the BPM plot to the doc file
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(plot_filename, width=Inches(5))

    # add the SPO2 Histogram to the doc file
    paragraph = doc.add_paragraph()
    his = paragraph.add_run()
    his.add_picture(hist_filename, width=Inches(5))

    # print that the doc file is generated in the console
    print(doc_name_print + " is generated")

    # Save the Word document
    doc.save(doc_name)

    # send word document to telegram
    bot = telebot.TeleBot("6140302269:AAG5rMISL5xamoIG5dnJcDuaJOK9qt1vWQU")
    chat_id = "639643824"
    doc = open(doc_name, 'rb')
    bot.send_document(chat_id, doc)
    print("Document sent successfully")


# Schedule the script to run every 1 sec
schedule.every(1).seconds.do(extract_data_and_generate_report)


def change_table_dimensions(table, width):
    # Change the table width
    table.width = width

    # Adjust the cell widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = width / len(row.cells)

# to keep the code running
while True:
    schedule.run_pending()
    time.sleep(1)
